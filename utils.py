# utils.py
from docx import Document
from docx2pdf import convert
import os

def put_name_to_pdf(input_docx_file, name_file, FIO, output_directory):
    doc = Document(input_docx_file)

    for table in doc.tables:
        for col in table.columns:
            for cell in col.cells:
                for paragraph in cell.paragraphs:
                    if '{{person_FIO}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{person_FIO}}', FIO)

    updated_docx_file = 'updated_doc.docx'
    doc.save(updated_docx_file)

    pdf_file = os.path.join(output_directory, f'{name_file}.pdf')
    convert(updated_docx_file, pdf_file)

    os.remove(updated_docx_file)


    
